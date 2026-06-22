from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Styling tuned for AI/ATS screening + a human 6-second skim (2026 research):
# single column, native text, standard headings, strong bold/size hierarchy,
# generous whitespace. No tables / columns / text boxes / images / header-footer
# content — all of which corrupt text extraction. One restrained dark accent.
_ACCENT = RGBColor(0x1A, 0x1A, 0x2E)  # dark navy — high contrast, parser-neutral
_BASE_FONT = "Calibri"                # ~99% parse-completeness across major ATS
_BODY_PT = 10.5


def build_resume_docx(data: dict) -> BytesIO:
    doc = Document()
    _set_base_font(doc)
    _set_margins(doc, top=0.75, bottom=0.75, left=0.8, right=0.8)

    contact = data.get("contact", {})
    _add_name_header(doc, contact.get("name", ""))
    if data.get("headline"):
        _add_headline(doc, data["headline"])
    _add_contact_line(doc, contact)

    _add_section_heading(doc, "SUMMARY")
    p = doc.add_paragraph(data.get("summary", ""))
    if p.runs:
        p.runs[0].font.size = Pt(_BODY_PT)

    _add_section_heading(doc, "TECHNICAL SKILLS")
    _add_skills(doc, data.get("skills", []))

    _add_section_heading(doc, "EXPERIENCE")
    for job in data.get("experience", []):
        _add_job_entry(doc, job)

    projects = data.get("projects", [])
    if projects:
        _add_section_heading(doc, "PROJECTS")
        for proj in projects:
            _add_project_entry(doc, proj)

    _add_section_heading(doc, "EDUCATION")
    for edu in data.get("education", []):
        _add_education_entry(doc, edu)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_cover_letter_docx(data: dict) -> BytesIO:
    doc = Document()
    _set_base_font(doc)
    _set_margins(doc, top=1.0, bottom=1.0, left=1.25, right=1.25)

    contact = data.get("contact", {})
    _add_name_header(doc, contact.get("name", ""))
    _add_contact_line(doc, contact)
    doc.add_paragraph("")

    cover_letter = data.get("cover_letter", "")
    for para_text in cover_letter.split("\n\n"):
        p = doc.add_paragraph(para_text.strip())
        p.paragraph_format.space_after = Pt(12)
        if p.runs:
            p.runs[0].font.size = Pt(11)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _set_base_font(doc):
    """Pin the Normal style so stray runs inherit a safe ATS font and uniform
    spacing, instead of Word's silent default — and so whitespace comes from
    paragraph spacing, not blank lines the parser reads as noise."""
    style = doc.styles["Normal"]
    style.font.name = _BASE_FONT
    style.font.size = Pt(_BODY_PT)
    pf = style.paragraph_format
    pf.line_spacing = 1.08
    pf.space_after = Pt(4)


def _set_margins(doc, top, bottom, left, right):
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def _add_name_header(doc, name: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(20)


def _add_headline(doc, headline: str):
    """Target-track line, top third of page 1 — the highest-leverage real estate
    for both a 6-second human skim and an LLM that weights early content."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(headline)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = _ACCENT


def _add_contact_line(doc, contact: dict):
    # Contact must live in the BODY — many parsers skip header/footer regions.
    parts = [contact.get("email", ""), contact.get("phone", ""),
             contact.get("location", ""), contact.get("links", "")]
    line = "  |  ".join(p for p in parts if p)
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p.runs:
        p.runs[0].font.size = Pt(10)


def _add_section_heading(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = _ACCENT
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A1A2E")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_skills(doc, skills: list):
    """Grouped, comma-separated plain text — one category per line so each
    token extracts as its own SKILL entity (a table would extract ZERO). If the
    model returns flat items (no 'Category:' prefix), fall back to a single
    pipe-joined line."""
    grouped = any(":" in s for s in skills)
    if not grouped:
        p = doc.add_paragraph(" | ".join(skills))
        if p.runs:
            p.runs[0].font.size = Pt(_BODY_PT)
        return
    for s in skills:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        if ":" in s:
            cat, items = s.split(":", 1)
            r1 = p.add_run(cat.strip() + ": ")
            r1.bold = True
            r1.font.size = Pt(_BODY_PT)
            r2 = p.add_run(items.strip())
            r2.font.size = Pt(_BODY_PT)
        else:
            r = p.add_run(s)
            r.font.size = Pt(_BODY_PT)


def _add_job_entry(doc, job: dict):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(job.get("company", ""))
    r.bold = True
    r.font.size = Pt(11)
    title_run = p.add_run(f"  —  {job.get('title', '')}")
    title_run.bold = True
    title_run.font.size = Pt(10.5)
    meta = doc.add_paragraph(f"{job.get('duration', '')}  |  {job.get('location', '')}")
    if meta.runs:
        meta.runs[0].font.size = Pt(9)
        meta.runs[0].italic = True
    meta.paragraph_format.space_after = Pt(2)
    for bullet in job.get("bullets", []):
        bp = doc.add_paragraph(bullet, style="List Bullet")
        if bp.runs:
            bp.runs[0].font.size = Pt(_BODY_PT)


def _add_project_entry(doc, proj: dict):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(proj.get("name", ""))
    r.bold = True
    r.font.size = Pt(_BODY_PT)
    desc = proj.get("description") or proj.get("detail")
    if desc:
        dr = p.add_run(f"  —  {desc}")
        dr.font.size = Pt(10)
    for bullet in proj.get("bullets", []):
        bp = doc.add_paragraph(bullet, style="List Bullet")
        if bp.runs:
            bp.runs[0].font.size = Pt(_BODY_PT)


def _add_education_entry(doc, edu: dict):
    p = doc.add_paragraph()
    r = p.add_run(edu.get("institution", ""))
    r.bold = True
    r.font.size = Pt(_BODY_PT)
    detail_run = p.add_run(f"  —  {edu.get('degree', '')}  ({edu.get('graduated', '')})")
    detail_run.font.size = Pt(_BODY_PT)
    for detail in edu.get("details", []):
        dp = doc.add_paragraph(detail, style="List Bullet")
        if dp.runs:
            dp.runs[0].font.size = Pt(_BODY_PT)
